
from __future__ import annotations

from datetime import date, datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _fmt(value: Any) -> str:
    if value in (None, ""):
        return "—"
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y %H:%M")
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    return str(value)


def _add_header(doc: Document, title: str = "ООО «ФармИнжиниринг»") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph()


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not rows:
        doc.add_paragraph("Нет данных.")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build_manager_report(
    path: Path,
    projects: list[dict],
    *,
    include_stages: bool,
    include_tasks: bool,
    include_materials: bool,
) -> None:
    doc = Document()
    _add_header(doc, "ООО «ФармИнжиниринг»")
    doc.add_heading("Отчёт по проектам", level=1)

    for project in projects:
        doc.add_heading(_fmt(project.get("name")), level=2)
        doc.add_paragraph(f"Адрес: {_fmt(project.get('address'))}")
        doc.add_paragraph(f"Статус: {_fmt(project.get('status'))}")
        doc.add_paragraph(
            f"Период: {_fmt(project.get('start_date'))} — {_fmt(project.get('end_date'))}"
        )

        if include_stages and project.get("stages"):
            doc.add_heading("Этапы", level=3)
            _add_table(
                doc,
                ["Этап", "Статус", "План начало", "План конец", "Бригада"],
                [
                    [
                        _fmt(s.get("stage_name")),
                        _fmt(s.get("stage_status")),
                        _fmt(s.get("planned_start")),
                        _fmt(s.get("planned_end")),
                        _fmt(s.get("brigade_name")),
                    ]
                    for s in project["stages"]
                ],
            )

        if include_tasks and project.get("tasks"):
            doc.add_heading("Задачи", level=3)
            _add_table(
                doc,
                ["Задача", "Этап", "Статус", "План начало", "План конец", "Факт начало", "Факт конец"],
                [
                    [
                        _fmt(t.get("task_name")),
                        _fmt(t.get("stage_name")),
                        _fmt(t.get("task_status")),
                        _fmt(t.get("planned_start")),
                        _fmt(t.get("planned_end")),
                        _fmt(t.get("actual_start")),
                        _fmt(t.get("actual_end")),
                    ]
                    for t in project["tasks"]
                ],
            )

        if include_materials and project.get("materials"):
            doc.add_heading("Материалы", level=3)
            _add_table(
                doc,
                ["Материал", "Этап", "План", "Факт"],
                [
                    [
                        _fmt(m.get("material_name")),
                        _fmt(m.get("stage_name")),
                        _fmt(m.get("planned_quantity")),
                        _fmt(m.get("actual_quantity")),
                    ]
                    for m in project["materials"]
                ],
            )
        doc.add_paragraph()

    doc.save(str(path))


def build_accountant_report(
    path: Path,
    rows: list[dict],
    *,
    date_from: date,
    date_to: date,
    employee_label: str,
) -> None:
    doc = Document()
    _add_header(doc, "ООО «ФармИнжиниринг»")
    doc.add_heading("Отчёт по отработанным часам", level=1)
    doc.add_paragraph(f"Период: {date_from.strftime('%d.%m.%Y')} — {date_to.strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"Сотрудники: {employee_label}")
    doc.add_paragraph()

    total = sum(float(r.get("hours") or 0) for r in rows)
    _add_table(
        doc,
        ["Сотрудник", "Проект", "Этап", "Задача", "Часы", "Дата"],
        [
            [
                _fmt(r.get("employee")),
                _fmt(r.get("project")),
                _fmt(r.get("stage")),
                _fmt(r.get("task")),
                _fmt(r.get("hours")),
                _fmt(r.get("work_date")),
            ]
            for r in rows
        ],
    )
    doc.add_paragraph(f"Итого часов: {total:.2f}")
    doc.save(str(path))
